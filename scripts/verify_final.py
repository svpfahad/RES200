"""Final verification script for the JURI submission docx."""
import sys
import zipfile
from pathlib import Path

import re
from docx import Document
from lxml import etree

sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(r"C:\Users\Fahad\Downloads\RES200")
DOCX = ROOT / "XGBoost_pKa_Prediction_REVISION_FINAL.docx"

print(f"Verifying: {DOCX}")
print(f"Size: {DOCX.stat().st_size / 1024:.1f} KB\n")

# ---------- Track-change residue check
with zipfile.ZipFile(DOCX) as z:
    body = z.read("word/document.xml").decode("utf-8")
    files = z.namelist()

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
root = etree.fromstring(body.encode("utf-8"))

residue = {
    "tracked insertions": len(list(root.iter(f"{W}ins"))),
    "tracked deletions":  len(list(root.iter(f"{W}del"))),
    "rPrChange":          len(list(root.iter(f"{W}rPrChange"))),
    "pPrChange":          len(list(root.iter(f"{W}pPrChange"))),
    "comment refs":       len(list(root.iter(f"{W}commentReference"))),
    "comment range":      len(list(root.iter(f"{W}commentRangeStart"))),
}
ok = all(v == 0 for v in residue.values())
print(f"[{'OK' if ok else 'FAIL'}] Track-change / comment residue:")
for k, v in residue.items():
    print(f"      {k}: {v}")

# ---------- Comment xml parts dropped
comment_parts = [f for f in files if "comments" in f.lower()]
print(f"\n[{'OK' if not comment_parts else 'FAIL'}] Comment XML parts:")
for c in comment_parts:
    print(f"      {c}")
if not comment_parts:
    print("      (none — good)")

# ---------- Document text checks
doc = Document(DOCX)
paras = doc.paragraphs

# Front matter
print("\n=== FRONT MATTER ===")
for i, p in enumerate(paras[:6]):
    if p.text.strip():
        print(f"P{i} [{p.style.name}]: {p.text}")

# Abstract word count
abstract_idx = next((i for i, p in enumerate(paras)
                     if p.text.startswith("The acid dissociation constant")), None)
if abstract_idx is not None:
    wc = len(paras[abstract_idx].text.split())
    print(f"\n[{'OK' if wc <= 250 else 'FAIL'}] Abstract word count: {wc} (must be ≤ 250)")

# Tables
print("\n=== TABLES ===")
for ti, t in enumerate(doc.tables):
    print(f"Table {ti+1}: {len(t.rows)}r × {len(t.columns)}c")
    if ti == 0:
        # Confirm Table 1 says 1,135 features
        cells = [c.text for r in t.rows for c in r.cells]
        if "1,135" in " ".join(cells):
            print("[OK] Table 1 contains '1,135 features'")
        else:
            print("[FAIL] Table 1 missing 1,135")

# References count
ref_lines = [p.text for p in paras
             if re.match(r"^\d{1,2}\.\s", p.text or "")]
print(f"\n[{'OK' if len(ref_lines) >= 15 else 'FAIL'}] References found: {len(ref_lines)}")
for r in ref_lines[:3]:
    print(f"      {r[:100]}…")

# Figure media presence
print("\n=== EMBEDDED FIGURES ===")
for n in sorted(files):
    if "media/image" in n:
        info = next(z.getinfo(n) for z in [zipfile.ZipFile(DOCX)])
        print(f"  {n}: {info.file_size:,} bytes")

# Author line check
author_text = paras[1].text if len(paras) > 1 else ""
fahad_star = "Fahad Atwi *" in author_text or "Fahad Atwi*" in author_text
khater_star = "Al-Khater *" in author_text or "Al-Khater*" in author_text
print(f"\n[{'OK' if not fahad_star else 'FAIL'}] Fahad's '*' removed (current: '{author_text}')")
print(f"[{'OK' if khater_star else 'FAIL'}] Al-Khater '*' present")

print("\n=== DONE ===")
