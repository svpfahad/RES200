from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt


PACKAGE_FILES = [
    "00_PACKAGE_README.md",
    "01_ATTORNEY_REVIEW_BRIEF.md",
    "02_SPECIFICATION_DRAFT.md",
    "03_CLAIMS_DRAFT.md",
    "04_DRAWING_PACKET.md",
    "05_PRIOR_ART_AND_IDS_CANDIDATES.md",
    "06_US_FILING_CHECKLIST.md",
    "07_SAIP_FILING_CHECKLIST.md",
    "08_EMAILS_TO_SEND.md",
    "09_INVENTOR_FACTS_NEEDED.md",
]


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        parts = [part.strip() for part in lines[index].strip().strip("|").split("|")]
        if parts and not all(set(part) <= {"-", ":", " "} for part in parts):
            rows.append(parts)
        index += 1
    return rows, index


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    for row_values in rows:
        cells = table.add_row().cells
        for column_index in range(column_count):
            cells[column_index].text = row_values[column_index] if column_index < len(row_values) else ""


def add_markdown(document: Document, text: str) -> None:
    lines = text.splitlines()
    in_code = False
    code_buffer: list[str] = []
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                run = paragraph.add_run("\n".join(code_buffer))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_buffer.append(raw)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif len(stripped) > 3 and stripped[0].isdigit() and ". " in stripped[:5]:
            document.add_paragraph(stripped.split(". ", 1)[1].strip(), style="List Number")
        else:
            document.add_paragraph(stripped)
        index += 1


def convert_one(source: Path, destination: Path) -> None:
    document = Document()
    add_markdown(document, source.read_text(encoding="utf-8"))
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def convert_combined(package_dir: Path, destination: Path) -> None:
    document = Document()
    for file_index, filename in enumerate(PACKAGE_FILES):
        source = package_dir / filename
        if not source.exists():
            continue
        if file_index > 0:
            document.add_page_break()
        add_markdown(document, source.read_text(encoding="utf-8"))
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export patent package Markdown files to DOCX.")
    parser.add_argument("--package-dir", default="patent_package")
    parser.add_argument("--out", default="patent_package_docx")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    package_dir = Path(args.package_dir)
    out_dir = Path(args.out)
    for filename in PACKAGE_FILES:
        source = package_dir / filename
        if source.exists():
            convert_one(source, out_dir / source.with_suffix(".docx").name)
    convert_combined(package_dir, out_dir / "COMBINED_PATENT_PACKAGE.docx")
    print(f"Exported DOCX package to: {out_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
